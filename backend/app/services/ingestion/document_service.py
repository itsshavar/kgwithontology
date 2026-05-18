import csv
import json
import re
from io import BytesIO, StringIO
from pathlib import Path
from uuid import uuid4
from xml.etree import ElementTree

from fastapi import UploadFile

from app.core.config import settings
from app.models.source_document import SourceDocument
from app.services.ingestion.text_chunker import chunk_text

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".html",
    ".xml",
    ".yaml",
    ".yml",
    ".ttl",
    ".rdf",
    ".owl",
    ".n3",
    ".nt",
    ".jsonld",
}


def _extract_pdf_text(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(file_bytes))
        texts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text:
                texts.append(page_text)
        return "\n\n".join(texts)
    except Exception:
        return ""


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
        return "\n".join(paragraphs)
    except Exception:
        return ""


def _extract_xlsx_text(file_bytes: bytes) -> str:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
        rows: list[str] = []
        for worksheet in workbook.worksheets:
            rows.append(f"# Sheet: {worksheet.title}")
            for row in worksheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    rows.append("\t".join(values))
        return "\n".join(rows)
    except Exception:
        return ""


def _normalize_structured_text(filename: str, file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    suffix = Path(filename).suffix.lower()
    if suffix == ".csv":
        try:
            rows = csv.reader(StringIO(text))
            return "\n".join("\t".join(cell for cell in row if cell) for row in rows)
        except Exception:
            return text
    if suffix in {".json", ".jsonld"}:
        try:
            return json.dumps(json.loads(text), indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            return text
    if suffix == ".xml":
        try:
            root = ElementTree.fromstring(text)
            values = [value.strip() for value in root.itertext() if value and value.strip()]
            return "\n".join(values) or text
        except ElementTree.ParseError:
            return text
    if suffix == ".rtf":
        return re.sub(r"[{}\\][a-zA-Z0-9* -]* ?", " ", text).replace("\\'", "")
    return text


def extract_text(filename: str, content_type: str | None, file_bytes: bytes) -> str:
    if not file_bytes:
        return ""

    lower_name = filename.lower()
    suffix = Path(lower_name).suffix
    if suffix == ".pdf" or content_type == "application/pdf":
        extracted = _extract_pdf_text(file_bytes)
        if extracted:
            return extracted
    if suffix == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        extracted = _extract_docx_text(file_bytes)
        if extracted:
            return extracted
    if suffix == ".xlsx" or content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        extracted = _extract_xlsx_text(file_bytes)
        if extracted:
            return extracted
    if suffix in TEXT_EXTENSIONS or suffix == ".rtf" or content_type and content_type.startswith("text/"):
        return _normalize_structured_text(filename, file_bytes)
    return file_bytes.decode("utf-8", errors="ignore")


def save_bytes(project_id: int, filename: str, content_type: str | None, file_bytes: bytes) -> SourceDocument:
    project_dir = settings.upload_dir / f"project_{project_id}"
    project_dir.mkdir(parents=True, exist_ok=True)

    safe_name = Path(filename or "uploaded_file").name
    stored_name = f"{uuid4().hex}_{safe_name}"
    storage_path = project_dir / stored_name
    storage_path.write_bytes(file_bytes)

    raw_text = extract_text(safe_name, content_type, file_bytes)
    chunks = chunk_text(raw_text)

    return SourceDocument(
        project_id=project_id,
        filename=safe_name,
        content_type=content_type,
        storage_path=str(storage_path.relative_to(settings.data_dir.parent)),
        raw_text=raw_text,
        chunk_count=len(chunks),
        status="processed" if raw_text is not None else "uploaded",
    )


async def save_upload(project_id: int, upload: UploadFile) -> SourceDocument:
    file_bytes = await upload.read()
    return save_bytes(project_id, upload.filename or "uploaded_file", upload.content_type, file_bytes)
